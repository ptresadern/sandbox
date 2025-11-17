"""Document processing for various file formats"""

import re
from pathlib import Path
from typing import List, Dict, Any
from abc import ABC, abstractmethod

from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup


class DocumentProcessor(ABC):
    """Abstract base class for document processors"""

    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the file"""
        pass

    @abstractmethod
    def extract_text(self, file_path: Path) -> str:
        """Extract text from the document"""
        pass


class TextProcessor(DocumentProcessor):
    """Process plain text files"""

    EXTENSIONS = {".txt", ".text", ".log"}

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.EXTENSIONS

    def extract_text(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()


class PDFProcessor(DocumentProcessor):
    """Process PDF files"""

    EXTENSIONS = {".pdf"}

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.EXTENSIONS

    def extract_text(self, file_path: Path) -> str:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)


class WordProcessor(DocumentProcessor):
    """Process Microsoft Word documents"""

    EXTENSIONS = {".docx", ".doc"}

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.EXTENSIONS

    def extract_text(self, file_path: Path) -> str:
        # Note: .doc (old format) requires additional libraries like antiword
        # We'll focus on .docx for now
        if file_path.suffix.lower() == ".doc":
            raise ValueError(
                "Old .doc format not supported. Please convert to .docx or use antiword."
            )

        doc = DocxDocument(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)


class MarkdownProcessor(DocumentProcessor):
    """Process Markdown files"""

    EXTENSIONS = {".md", ".markdown", ".mdown"}

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.EXTENSIONS

    def extract_text(self, file_path: Path) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        # Convert markdown to HTML
        html = markdown.markdown(md_content)

        # Extract text from HTML
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n", strip=True)


class HTMLProcessor(DocumentProcessor):
    """Process HTML files"""

    EXTENSIONS = {".html", ".htm"}

    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.EXTENSIONS

    def extract_text(self, file_path: Path) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator="\n", strip=True)

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)

        return text


class DocumentManager:
    """Manage document processing with multiple processors"""

    def __init__(self):
        self.processors: List[DocumentProcessor] = [
            TextProcessor(),
            PDFProcessor(),
            WordProcessor(),
            MarkdownProcessor(),
            HTMLProcessor(),
        ]

    def can_process(self, file_path: Path) -> bool:
        """Check if any processor can handle this file"""
        return any(p.can_process(file_path) for p in self.processors)

    def extract_text(self, file_path: Path) -> str:
        """Extract text from a document"""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor.extract_text(file_path)
        raise ValueError(f"No processor found for file: {file_path}")

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks based on approximate token count"""
        # Simple word-based chunking (approximate)
        # More sophisticated tokenization could use tiktoken or similar
        words = text.split()

        chunks = []
        start = 0

        while start < len(words):
            end = start + chunk_size
            chunk_words = words[start:end]
            chunks.append(" ".join(chunk_words))

            # Move start position with overlap
            start = end - overlap

        return chunks

    def process_document(
        self, file_path: Path, chunk_size: int = 500, overlap: int = 50
    ) -> Dict[str, Any]:
        """Process a document and return chunks with metadata"""
        try:
            text = self.extract_text(file_path)
            chunks = self.chunk_text(text, chunk_size, overlap)

            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "text": text,
                "chunks": chunks,
                "num_chunks": len(chunks),
                "success": True,
                "error": None,
            }
        except Exception as e:
            return {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "text": "",
                "chunks": [],
                "num_chunks": 0,
                "success": False,
                "error": str(e),
            }

    def process_directory(
        self, directory: Path, chunk_size: int = 500, overlap: int = 50
    ) -> List[Dict[str, Any]]:
        """Process all supported documents in a directory"""
        results = []

        if not directory.exists():
            return results

        for file_path in directory.rglob("*"):
            if file_path.is_file() and self.can_process(file_path):
                result = self.process_document(file_path, chunk_size, overlap)
                results.append(result)

        return results


# Global document manager instance
doc_manager = DocumentManager()
